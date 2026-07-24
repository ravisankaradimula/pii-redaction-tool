# redactor.py
from __future__ import annotations

import re
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Tuple
from uuid import uuid4

from docx import Document
from faker import Faker
import spacy

INPUT_DIR = Path("input_docs")
OUTPUT_DIR = Path("output_docs")
OUTPUT_DIR.mkdir(exist_ok=True)

fake = Faker("en_US")
ENTITY_MAP: Dict[str, str] = {}

# Regex patterns for structured PII
EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(
    r"\b(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)\d{3}[-.\s]?\d{4}\b"
)
CREDIT_CARD_RE = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")
IPV4_RE = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")
DOB_RE = re.compile(
    r"\b(?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])[-/](?:19|20)\d{2}\b"
    r"|"
    r"\b(?:19|20)\d{2}[-/](?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])\b"
)
SSN_PAN_RE = re.compile(
    r"\b(?:\d{3}-\d{2}-\d{4}|\d{4}[- ]\d{4}[- ]\d{4}|\d{4}[- ]\d{6}|\d{9,12})\b"
)
ADDRESS_RE = re.compile(
    r"\b\d{1,5}\s+[A-Za-z0-9.\- ]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|Lane|Ln|Parkway|Pkwy|Way|Court|Ct|Circle|Cir|Terrace|Ter|Place|Pl|Highway|Hwy|Trail|Trl)\b",
    flags=re.IGNORECASE,
)

try:
    NLP = spacy.load("en_core_web_sm")
except OSError as exc:
    raise SystemExit(
        "spaCy model 'en_core_web_sm' is not installed. "
        "Please run: python -m spacy download en_core_web_sm"
    ) from exc


def normalize_pii(value: str) -> str:
    return re.sub(r"\s+", "", value.strip().casefold())


def get_or_create_fake_value(kind: str, original: str) -> str:
    key = normalize_pii(original)
    if key in ENTITY_MAP:
        return ENTITY_MAP[key]

    if kind == "email":
        replacement = fake.ascii_free_email()
    elif kind == "phone":
        replacement = fake.numerify(text="+1-###-###-####")
    elif kind == "credit_card":
        replacement = fake.credit_card_number()
    elif kind == "ip":
        replacement = fake.ipv4_public()
    elif kind == "dob":
        replacement = fake.date_of_birth(minimum_age=18, maximum_age=75).strftime("%m/%d/%Y")
    elif kind == "id":
        lowered = original.lower()
        if "ssn" in lowered or "social" in lowered or "-" in original:
            replacement = fake.ssn()
        else:
            replacement = fake.bothify(text="#### #### ####")
    elif kind == "address":
        replacement = fake.street_address()
    elif kind == "person":
        replacement = fake.name()
    elif kind == "org":
        replacement = fake.company()
    elif kind == "location":
        replacement = fake.city()
    else:
        replacement = original

    ENTITY_MAP[key] = replacement
    return replacement


def replace_structured_matches(text: str, pattern: re.Pattern, kind: str) -> str:
    matches = list(pattern.finditer(text))
    if not matches:
        return text

    for match in reversed(matches):
        original = match.group(0)
        replacement = get_or_create_fake_value(kind, original)
        text = f"{text[:match.start()]}{replacement}{text[match.end():]}"
    return text


def redact_structured_entities(text: str) -> str:
    text = replace_structured_matches(text, EMAIL_RE, "email")
    text = replace_structured_matches(text, PHONE_RE, "phone")
    text = replace_structured_matches(text, CREDIT_CARD_RE, "credit_card")
    text = replace_structured_matches(text, IPV4_RE, "ip")
    text = replace_structured_matches(text, DOB_RE, "dob")
    text = replace_structured_matches(text, SSN_PAN_RE, "id")
    text = replace_structured_matches(text, ADDRESS_RE, "address")
    return text


def redact_named_entities(text: str) -> str:
    doc = NLP(text)
    spans: List[Tuple[int, int, str, str]] = []

    for ent in doc.ents:
        if ent.label_ in {"PERSON", "ORG", "GPE", "LOC", "FAC"}:
            spans.append((ent.start_char, ent.end_char, ent.label_, ent.text))

    for start, end, label, original in reversed(spans):
        kind = {
            "PERSON": "person",
            "ORG": "org",
            "GPE": "location",
            "LOC": "location",
            "FAC": "address",
        }.get(label, "location")
        replacement = get_or_create_fake_value(kind, original)
        text = f"{text[:start]}{replacement}{text[end:]}"

    return text


def redact_text(text: str) -> str:
    text = redact_structured_entities(text)
    text = redact_named_entities(text)
    return text


def sanitize_paragraph(paragraph) -> None:
    sanitized_text = redact_text(paragraph.text)
    if sanitized_text != paragraph.text:
        paragraph.text = sanitized_text


def sanitize_document(document: Document) -> None:
    for paragraph in document.paragraphs:
        sanitize_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    sanitize_paragraph(paragraph)


def get_input_document_path() -> Path:
    candidates = sorted(INPUT_DIR.glob("*.docx"))
    if not candidates:
        raise FileNotFoundError(f"No .docx files found in {INPUT_DIR}")
    return candidates[0]


def print_evaluation_summary() -> None:
    from evaluation import calculate_metrics, format_percentage, run_simulation

    tp, fp, fn, tn = run_simulation()
    recall, precision, accuracy = calculate_metrics(tp, fp, fn, tn)

    print("\nPII Redaction Evaluation Metrics")
    print("-" * 40)
    print(f"Recall: {format_percentage(recall)}")
    print(f"Precision: {format_percentage(precision)}")
    print(f"Accuracy: {format_percentage(accuracy)}")
    print("-" * 40)
    print(f"TP={tp}  FP={fp}  FN={fn}  TN={tn}")


def main() -> None:
    input_file = get_input_document_path()
    output_file = OUTPUT_DIR / f"{input_file.stem}_sanitized_{uuid4().hex[:8]}.docx"

    document = Document(input_file)
    sanitize_document(document)
    document.save(output_file)

    print(f"Sanitized document saved to: {output_file}")
    print_evaluation_summary()


if __name__ == "__main__":
    main()